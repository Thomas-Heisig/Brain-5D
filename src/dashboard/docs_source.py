"""Safe read-only access to repository documentation and data files.

Supports: .md, .txt, .docx, .xlsx, .csv, .json, .pdf (metadata only)
"""

from __future__ import annotations

import csv
import json
from dataclasses import dataclass
from datetime import datetime
from enum import Enum
from functools import lru_cache
from pathlib import Path
from typing import Any, Optional

from .models import JSONValue

# Optional imports with proper fallbacks – using lower-case names to avoid Pylance constant redefinition warnings.
try:
    from docx import Document as DocxDocument  # type: ignore[import-untyped]

    has_docx: bool = True
except ImportError:
    has_docx = False
    DocxDocument = None  # type: ignore[assignment]

try:
    from openpyxl import load_workbook  # type: ignore[import-untyped]

    has_openpyxl: bool = True
except ImportError:
    has_openpyxl = False
    load_workbook = None  # type: ignore[assignment]

try:
    import PyPDF2  # type: ignore[import-untyped]

    has_pypdf2: bool = True
except ImportError:
    has_pypdf2 = False
    PyPDF2 = None  # type: ignore[assignment]

__all__ = [
    "FileType",
    "DocumentationEntry",
    "DocumentationSource",
    "create_docs_source",
]


class FileType(Enum):
    """Supported document file types."""

    MARKDOWN = "markdown"
    TEXT = "text"
    DOCX = "docx"
    XLSX = "xlsx"
    CSV = "csv"
    JSON = "json"
    PDF = "pdf"
    UNKNOWN = "unknown"

    @classmethod
    def from_extension(cls, ext: str) -> FileType:
        """Map file extension to FileType."""
        ext = ext.lower().lstrip(".")
        mapping: dict[str, FileType] = {
            "md": cls.MARKDOWN,
            "markdown": cls.MARKDOWN,
            "txt": cls.TEXT,
            "text": cls.TEXT,
            "docx": cls.DOCX,
            "xlsx": cls.XLSX,
            "xls": cls.XLSX,
            "csv": cls.CSV,
            "json": cls.JSON,
            "pdf": cls.PDF,
        }
        return mapping.get(ext, cls.UNKNOWN)

    @property
    def is_editable(self) -> bool:
        """Whether this file type can be edited in a text editor."""
        return self in {FileType.MARKDOWN, FileType.TEXT, FileType.CSV, FileType.JSON}

    @property
    def is_binary(self) -> bool:
        """Whether this file type is binary."""
        return self in {FileType.DOCX, FileType.XLSX, FileType.PDF}


@dataclass(frozen=True, slots=True)
class DocumentationEntry:
    """Rich metadata for a documentation or data file."""

    name: str
    path: str
    size_bytes: int
    file_type: FileType
    modified_time: str
    content_preview: Optional[str] = None
    word_count: Optional[int] = None
    line_count: Optional[int] = None
    sheet_names: Optional[tuple[str, ...]] = None
    supported: bool = True

    def to_json(self) -> dict[str, JSONValue]:
        """Return JSON-ready document metadata."""
        result: dict[str, JSONValue] = {
            "name": self.name,
            "path": self.path,
            "size_bytes": self.size_bytes,
            "file_type": self.file_type.value,
            "modified_time": self.modified_time,
            "supported": self.supported,
        }
        if self.content_preview is not None:
            result["content_preview"] = self.content_preview
        if self.word_count is not None:
            result["word_count"] = self.word_count
        if self.line_count is not None:
            result["line_count"] = self.line_count
        if self.sheet_names is not None:
            result["sheet_names"] = list(self.sheet_names)
        return result


class DocumentationSource:
    """Expose documentation and data files with multi-format support.

    Provides safe, read-only access to files below a fixed root directory.
    Supports content extraction for .md, .txt, .docx, .xlsx, .csv, .json.
    """

    def __init__(
        self,
        docs_root: Path,
        max_preview_chars: int = 500,
        max_file_size_mb: int = 50,
        enable_caching: bool = True,
    ) -> None:
        """Initialize documentation source.

        Args:
            docs_root: Root directory for documentation files.
            max_preview_chars: Maximum characters for content preview.
            max_file_size_mb: Maximum file size to process (MB).
            enable_caching: Whether to cache file listings.
        """
        self.docs_root = docs_root.resolve()
        self.max_preview_chars = max_preview_chars
        self.max_file_size_bytes = max_file_size_mb * 1024 * 1024
        self.enable_caching = enable_caching

    def list_documents(self, recursive: bool = False) -> tuple[DocumentationEntry, ...]:
        """List all supported documents in stable order.

        Args:
            recursive: Whether to scan subdirectories recursively.

        Returns:
            Tuple of DocumentationEntry objects.
        """
        if not self.docs_root.is_dir():
            return ()

        pattern = "**/*" if recursive else "*"
        entries: list[DocumentationEntry] = []

        for path in sorted(self.docs_root.glob(pattern)):
            if not path.is_file():
                continue
            if self._is_excluded(path):
                continue

            entry = self._build_entry(path)
            if entry is not None:
                entries.append(entry)

        return tuple(entries)

    def list_documents_by_type(self, file_type: FileType) -> tuple[DocumentationEntry, ...]:
        """List documents filtered by file type."""
        return tuple(
            entry for entry in self.list_documents(recursive=True) if entry.file_type == file_type
        )

    def get_document(self, path: str) -> DocumentationEntry:
        """Get metadata for a specific document."""
        resolved_path = self._resolve_path(path)
        if resolved_path is None or not resolved_path.is_file():
            raise FileNotFoundError(f"Document not found: {path}")

        entry = self._build_entry(resolved_path)
        if entry is None:
            raise ValueError(f"Unsupported file type: {path}")
        return entry

    def read_content(self, path: str) -> str:
        """Read the full content of a document (text extraction)."""
        resolved_path = self._resolve_path(path)
        if resolved_path is None or not resolved_path.is_file():
            raise FileNotFoundError(f"Document not found: {path}")

        file_type = FileType.from_extension(resolved_path.suffix)
        return self._extract_content(resolved_path, file_type, full=True)

    def read_preview(self, path: str) -> str:
        """Read a preview of the document content."""
        resolved_path = self._resolve_path(path)
        if resolved_path is None or not resolved_path.is_file():
            raise FileNotFoundError(f"Document not found: {path}")

        file_type = FileType.from_extension(resolved_path.suffix)
        return self._extract_content(resolved_path, file_type, full=False)

    def get_directory_structure(self) -> dict[str, Any]:
        """Get the full directory tree with metadata."""
        if not self.docs_root.is_dir():
            return {"path": str(self.docs_root), "children": []}

        return self._build_tree(self.docs_root)

    def _build_tree(self, path: Path, relative_path: str = "") -> dict[str, Any]:
        """Recursively build directory tree."""
        result: dict[str, Any] = {
            "name": path.name if relative_path else str(path),
            "path": relative_path or ".",
            "type": "directory",
            "children": [],
        }

        for child in sorted(path.iterdir()):
            if self._is_excluded(child):
                continue

            if child.is_dir():
                child_path = str(child.relative_to(self.docs_root))
                result["children"].append(self._build_tree(child, child_path))
            elif child.is_file():
                entry = self._build_entry(child)
                if entry is not None:
                    result["children"].append(
                        {
                            "name": child.name,
                            "path": str(child.relative_to(self.docs_root)),
                            "type": "file",
                            "size_bytes": child.stat().st_size,
                            "file_type": entry.file_type.value,
                            "modified_time": entry.modified_time,
                        }
                    )

        return result

    def _build_entry(self, path: Path) -> Optional[DocumentationEntry]:
        """Build a DocumentationEntry from a file path."""
        try:
            stat = path.stat()
            file_type = FileType.from_extension(path.suffix)
            supported = self._is_supported(file_type)

            # Base metadata
            name = path.name
            rel_path = str(path.relative_to(self.docs_root))
            size_bytes = stat.st_size
            modified_time = datetime.fromtimestamp(stat.st_mtime).isoformat()

            entry_kwargs: dict[str, Any] = {
                "name": name,
                "path": rel_path,
                "size_bytes": size_bytes,
                "file_type": file_type,
                "modified_time": modified_time,
                "supported": supported,
            }

            # Skip content extraction for unsupported or large files
            if not supported or size_bytes > self.max_file_size_bytes:
                return DocumentationEntry(**entry_kwargs)  # type: ignore[arg-type]

            # Extract content preview and metrics
            try:
                content = self._extract_content(path, file_type, full=False)
                if content:
                    entry_kwargs["content_preview"] = content[: self.max_preview_chars]

                    if file_type.is_editable:
                        full_content = self._extract_content(path, file_type, full=True)
                        if full_content:
                            entry_kwargs["word_count"] = len(full_content.split())
                            entry_kwargs["line_count"] = full_content.count("\n") + 1

                # Sheet names for Excel
                if file_type == FileType.XLSX and has_openpyxl and load_workbook is not None:
                    try:
                        wb = load_workbook(path, read_only=True, data_only=True)
                        entry_kwargs["sheet_names"] = tuple(wb.sheetnames)
                        wb.close()
                    except Exception:
                        pass

            except Exception:
                pass

            return DocumentationEntry(**entry_kwargs)  # type: ignore[arg-type]

        except Exception:
            return None

    def _extract_content(self, path: Path, file_type: FileType, full: bool = False) -> str:
        """Extract text content from a file based on its type."""
        if file_type == FileType.MARKDOWN or file_type == FileType.TEXT:
            return path.read_text(encoding="utf-8", errors="replace")

        if file_type == FileType.CSV:
            return self._extract_csv_content(path, full)

        if file_type == FileType.JSON:
            return self._extract_json_content(path, full)

        if file_type == FileType.DOCX and has_docx and DocxDocument is not None:
            return self._extract_docx_content(path, full)

        if file_type == FileType.XLSX and has_openpyxl and load_workbook is not None:
            return self._extract_xlsx_content(path, full)

        if file_type == FileType.PDF and has_pypdf2 and PyPDF2 is not None:
            return self._extract_pdf_content(path, full)

        return f"Content extraction not available for {file_type.value} files."

    def _extract_csv_content(self, path: Path, full: bool) -> str:
        """Extract CSV content as readable text."""
        try:
            lines: list[str] = []
            with open(path, "r", encoding="utf-8", errors="replace") as f:
                reader = csv.reader(f)
                if not full:
                    for i, row in enumerate(reader):
                        if i >= 10:
                            lines.append("... (truncated)")
                            break
                        lines.append(" | ".join(row))
                else:
                    for row in reader:
                        lines.append(" | ".join(row))
            return "\n".join(lines)
        except Exception:
            return f"[Could not parse CSV: {path.name}]"

    def _extract_json_content(self, path: Path, full: bool) -> str:
        """Extract JSON content as formatted text."""
        try:
            data = json.loads(path.read_text(encoding="utf-8", errors="replace"))
            if not full:
                if isinstance(data, list):
                    # Type: ignore - JSON data has arbitrary structure that Pylance cannot infer
                    preview = data[:2]  # type: ignore[var-annotated]
                    if len(data) > 2:  # type: ignore[arg-type]
                        preview.append("...")  # type: ignore[attr-defined]
                    return json.dumps(preview, indent=2, ensure_ascii=False)
                elif isinstance(data, dict):
                    items = list(data.items())  # type: ignore[var-annotated, arg-type]
                    preview_dict = dict(items[:5])  # type: ignore[var-annotated, arg-type]
                    if len(items) > 5:  # type: ignore[arg-type]
                        preview_dict["..."] = f"({len(items) - 5} more keys)"  # type: ignore[index]
                    return json.dumps(preview_dict, indent=2, ensure_ascii=False)
            return json.dumps(data, indent=2, ensure_ascii=False)
        except Exception:
            return f"[Could not parse JSON: {path.name}]"

    def _extract_docx_content(self, path: Path, full: bool) -> str:
        """Extract text from DOCX file."""
        try:
            if DocxDocument is None:
                return "[python-docx not available]"

            doc = DocxDocument(str(path))
            paragraphs: list[str] = []
            for p in doc.paragraphs:
                if p.text.strip():
                    paragraphs.append(p.text)

            if not full:
                paragraphs = paragraphs[:20]
                if len(doc.paragraphs) > 20:
                    paragraphs.append("... (truncated)")
            return "\n".join(paragraphs)
        except Exception:
            return f"[Could not read DOCX: {path.name}]"

    def _extract_xlsx_content(self, path: Path, full: bool) -> str:
        """Extract text from XLSX file."""
        try:
            if load_workbook is None:
                return "[openpyxl not available]"

            wb = load_workbook(path, read_only=True, data_only=True)
            lines: list[str] = []
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                lines.append(f"\n=== Sheet: {sheet_name} ===\n")

                max_rows = 50 if not full else 1000
                row_count = 0
                for row in sheet.iter_rows(values_only=True):
                    if row_count >= max_rows:
                        lines.append("... (truncated)")
                        break
                    row_str = " | ".join(str(cell) if cell is not None else "" for cell in row)
                    if row_str.strip():
                        lines.append(row_str)
                    row_count += 1

            wb.close()
            return "\n".join(lines)
        except Exception:
            return f"[Could not read XLSX: {path.name}]"

    def _extract_pdf_content(self, path: Path, full: bool) -> str:
        """Extract text from PDF file."""
        try:
            if PyPDF2 is None:
                return "[PyPDF2 not available]"

            text_parts: list[str] = []
            with open(path, "rb") as f:
                pdf = PyPDF2.PdfReader(f)
                max_pages = 5 if not full else len(pdf.pages)
                for i in range(min(max_pages, len(pdf.pages))):
                    page = pdf.pages[i]
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                if not full and len(pdf.pages) > 5:
                    text_parts.append("... (truncated)")
            return "\n".join(text_parts)
        except Exception:
            return f"[Could not read PDF: {path.name}]"

    def _resolve_path(self, path: str) -> Optional[Path]:
        """Resolve a path relative to docs_root with security checks."""
        # Basic security: prevent path traversal
        if ".." in path or path.startswith("/") or path.startswith("\\"):
            return None

        try:
            full_path = (self.docs_root / path).resolve()
        except ValueError:
            return None

        # Ensure path is within docs_root
        try:
            full_path.relative_to(self.docs_root)
        except ValueError:
            return None

        return full_path

    def _is_excluded(self, path: Path) -> bool:
        """Check if a path should be excluded."""
        if path.name.startswith("."):
            return True
        exclude_dirs = {"__pycache__", ".git", ".venv", "node_modules", ".pytest_cache"}
        if path.name in exclude_dirs:
            return True
        return False

    def _is_supported(self, file_type: FileType) -> bool:
        """Check if a file type is supported for content extraction."""
        if file_type in {FileType.MARKDOWN, FileType.TEXT, FileType.CSV, FileType.JSON}:
            return True
        if file_type == FileType.DOCX:
            return has_docx
        if file_type == FileType.XLSX:
            return has_openpyxl
        if file_type == FileType.PDF:
            return has_pypdf2
        return False

    @lru_cache(maxsize=128)
    def _get_cached_list(self, recursive: bool) -> tuple[DocumentationEntry, ...]:
        """Cached version of list_documents for performance."""
        return self.list_documents(recursive)

    def invalidate_cache(self) -> None:
        """Invalidate the internal cache."""
        self._get_cached_list.cache_clear()

    def search_documents(
        self, query: str, file_types: Optional[list[FileType]] = None
    ) -> tuple[DocumentationEntry, ...]:
        """Search documents by filename (basic search)."""
        query_lower = query.lower()
        results: list[DocumentationEntry] = []
        for entry in self.list_documents(recursive=True):
            if file_types is not None and entry.file_type not in file_types:
                continue
            if query_lower in entry.name.lower():
                results.append(entry)
        return tuple(results)


def create_docs_source(docs_root: str | Path) -> DocumentationSource:
    """Factory function for creating a DocumentationSource."""
    if isinstance(docs_root, str):
        docs_root = Path(docs_root)
    return DocumentationSource(docs_root)