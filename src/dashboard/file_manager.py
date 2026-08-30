"""Unified file manager for Research and Docs sources in the dashboard.

Provides a unified API over both ``research/`` and ``docs/`` directories
with proper MIME type detection, binary file serving, and directory trees.
"""

from __future__ import annotations

import re
import subprocess
from http import HTTPStatus
from io import BytesIO
from pathlib import Path
from typing import Any, Iterator, cast
from urllib.parse import unquote

from .docs_source import DocumentationSource, create_docs_source
from .research_source import ResearchSource

# ---------------------------------------------------------------------------
# MIME type map
# ---------------------------------------------------------------------------

_MEDIA_TYPES: dict[str, str] = {
    ".html": "text/html; charset=utf-8",
    ".css": "text/css; charset=utf-8",
    ".js": "text/javascript; charset=utf-8",
    ".svg": "image/svg+xml",
    ".ico": "image/x-icon",
    ".png": "image/png",
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
    ".gif": "image/gif",
    ".webp": "image/webp",
    ".bmp": "image/bmp",
    ".mp4": "video/mp4",
    ".webm": "video/webm",
    ".ogg": "video/ogg",
    ".mov": "video/quicktime",
    ".avi": "video/x-msvideo",
    ".mp3": "audio/mpeg",
    ".wav": "audio/wav",
    ".flac": "audio/flac",
    ".aac": "audio/aac",
    ".m4a": "audio/mp4",
    ".opus": "audio/opus",
    ".pdf": "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
}

_BINARY_EXTENSIONS = frozenset({
    ".png", ".jpg", ".jpeg", ".gif", ".webp", ".svg",
    ".mp4", ".webm", ".ogg", ".mov", ".avi",
    ".mp3", ".wav", ".flac", ".aac", ".m4a", ".opus",
    ".pdf", ".docx", ".xlsx", ".xls",
    ".ico", ".bmp",
})


def _mime_type(ext: str) -> str:
    return _MEDIA_TYPES.get(ext.lower(), "application/octet-stream")


# ---------------------------------------------------------------------------
# Exceptions
# ---------------------------------------------------------------------------


class FileManagerError(Exception):
    """Base error for file manager operations."""


class SourceNotAvailableError(FileManagerError):
    """Raised when a requested source is not configured."""


class InvalidSourceError(FileManagerError):
    """Raised when an unknown source name is given."""


class PathTraversalError(FileManagerError):
    """Raised on path traversal attempts."""


# ---------------------------------------------------------------------------
# Markdown export helpers
# ---------------------------------------------------------------------------


def _escape_html(value: str) -> str:
    return (
        value.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
    )


def _markdown_to_html(content: str, title: str = "Document") -> str:
    """Convert a simple Markdown document to a standalone HTML page."""
    lines = content.splitlines()
    html_parts: list[str] = []
    in_code = False
    code_lines: list[str] = []

    for raw_line in lines:
        line = raw_line.rstrip()
        if line.startswith("```"):
            if in_code:
                html_parts.append("<pre><code>" + _escape_html("\n".join(code_lines)) + "</code></pre>")
                code_lines = []
            in_code = not in_code
            continue
        if in_code:
            code_lines.append(raw_line)
            continue

        # Heading
        match = re.match(r"^(#{1,6})\s+(.*)$", line)
        if match:
            level = len(match.group(1))
            html_parts.append(f"<h{level}>{_escape_html(match.group(2))}</h{level}>")
            continue

        # Unordered list
        if re.match(r"^\s*[-*+]\s+", line):
            text = re.sub(r"^\s*[-*+]\s+", "", line)
            html_parts.append(f"<li>{_escape_html(text)}</li>")
            continue

        # Ordered list
        match = re.match(r"^(\s*)\d+\.\s+(.*)$", line)
        if match:
            html_parts.append(f"<li>{_escape_html(match.group(2))}</li>")
            continue

        # Empty line
        if not line.strip():
            html_parts.append("<br>")
            continue

        # Paragraph with inline formatting
        para = _inline_markdown(line)
        html_parts.append(f"<p>{para}</p>")

    # Wrap loose list items in <ul>
    result: list[str] = []
    list_buffer: list[str] = []
    for part in html_parts:
        if part.startswith("<li>"):
            list_buffer.append(part)
        else:
            if list_buffer:
                result.append("<ul>" + "".join(list_buffer) + "</ul>")
                list_buffer = []
            result.append(part)
    if list_buffer:
        result.append("<ul>" + "".join(list_buffer) + "</ul>")

    body = "\n".join(result)
    return f"""<!DOCTYPE html>
<html lang=\"en\">
<head>
<meta charset=\"utf-8\">
<title>{_escape_html(title)}</title>
<style>
body {{ font-family: system-ui, -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif; max-width: 720px; margin: 2rem auto; padding: 0 1rem; line-height: 1.6; color: #1e293b; }}
h1, h2, h3 {{ color: #0f172a; }}
code {{ background: #f1f5f9; padding: 0.15rem 0.3rem; border-radius: 4px; }}
pre {{ background: #f1f5f9; padding: 0.75rem; border-radius: 8px; overflow-x: auto; }}
ul {{ padding-left: 1.25rem; }}
li {{ margin: 0.25rem 0; }}
</style>
</head>
<body>
{body}
</body>
</html>"""


def _inline_markdown(line: str) -> str:
    """Apply inline Markdown formatting (bold, italic, code links)."""
    # Use placeholders to protect markup during escaping
    placeholders: dict[str, str] = {}
    counter = 0

    def _protect(pattern: str, text: str) -> str:
        nonlocal counter
        out = text
        for m in re.finditer(pattern, text):
            key = f"\x00PLACEHOLDER{counter}\x00"
            counter += 1
            placeholders[key] = m.group(0)
            out = out.replace(m.group(0), key, 1)
        return out

    protected = _protect(r"`[^`]+`", line)
    protected = _protect(r"\*\*[^*]+\*\*", protected)
    protected = _protect(r"__[^_]+__", protected)
    protected = _protect(r"\*[^*]+\*", protected)
    protected = _protect(r"_[^_]+_", protected)
    protected = _protect(r"\[[^\]]+\]\([^)]+\)", protected)

    escaped = _escape_html(protected)
    for key, original in placeholders.items():
        # Convert protected Markdown to HTML
        html = original
        html = re.sub(r"`([^`]+)`", r"<code>\1</code>", html)
        html = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", html)
        html = re.sub(r"__(.+?)__", r"<strong>\1</strong>", html)
        html = re.sub(r"\*(.+?)\*", r"<em>\1</em>", html)
        html = re.sub(r"_(.+?)_", r"<em>\1</em>", html)
        html = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r'<a href="\2">\1</a>', html)
        escaped = escaped.replace(key, html)
    return escaped


def _iter_markdown_blocks(content: str) -> Iterator[dict[str, Any]]:
    """Yield Markdown blocks for DOCX export."""
    lines = content.splitlines()
    in_code = False
    code_lines: list[str] = []
    for raw_line in lines:
        line = raw_line.rstrip()
        if line.startswith("```"):
            if in_code:
                yield {"type": "code", "text": "\n".join(code_lines)}
                code_lines = []
            in_code = not in_code
            continue
        if in_code:
            code_lines.append(raw_line)
            continue

        match = re.match(r"^(#{1,6})\s+(.*)$", line)
        if match:
            yield {"type": "heading", "level": len(match.group(1)), "text": match.group(2)}
            continue

        if re.match(r"^\s*[-*+\d]\.\s+", line):
            text = re.sub(r"^\s*[-*+\d]\.\s+", "", line)
            yield {"type": "list", "text": text}
            continue

        if not line.strip():
            continue

        yield {"type": "paragraph", "text": line}


# ---------------------------------------------------------------------------
# File Manager
# ---------------------------------------------------------------------------


class FileManager:
    """Unified file manager for Research and Docs sources.

    Provides directory tree listing, file search, statistics, and
    content retrieval (text as JSON, binary as raw bytes).
    """

    def __init__(
        self,
        research_source: ResearchSource | None,
        docs_source: DocumentationSource | None,
        default_docs_root: Path,
    ) -> None:
        self._research = research_source
        self._docs = docs_source
        self._default_docs_root = default_docs_root

    # ------------------------------------------------------------------
    # Source resolution
    # ------------------------------------------------------------------

    def _resolve_source(self, source: str) -> tuple[Any, Path, str]:
        if source == "research":
            s = self._research
            if s is None or not s.is_available():
                raise SourceNotAvailableError("Research source is not configured.")
            return s, s.root(), "research"

        if source == "docs":
            s = self._docs or create_docs_source(self._default_docs_root)
            return s, s.docs_root, "docs"

        raise InvalidSourceError(f"Unknown source: {source}")

    def _root(self, source: str) -> Path:
        """Return the filesystem root for a source."""
        return self._resolve_source(source)[1]

    # ------------------------------------------------------------------
    # Directory tree
    # ------------------------------------------------------------------

    def get_tree(self, source: str = "research") -> dict[str, Any]:
        """Return the full directory tree for a source."""
        try:
            _obj, root, _ = self._resolve_source(source)
        except FileManagerError as exc:
            return {"available": False, "error": str(exc), "children": []}

        def _build(path: Path, rel_prefix: str = "") -> dict[str, Any]:
            children: list[dict[str, Any]] = []
            try:
                entries = sorted(path.iterdir())
            except PermissionError:
                return {
                    "name": path.name,
                    "path": rel_prefix or ".",
                    "type": "directory",
                    "children": [],
                }

            for child in entries:
                if child.name.startswith("."):
                    continue
                if child.name in {"__pycache__", ".git", ".venv", "node_modules"}:
                    continue

                if child.is_dir():
                    child_rel = (
                        str(child.relative_to(root)) if child != root else ""
                    )
                    children.append(_build(child, child_rel))
                elif child.is_file():
                    ext = child.suffix.lower()
                    children.append({
                        "name": child.name,
                        "path": str(child.relative_to(root)),
                        "type": "file",
                        "size_bytes": child.stat().st_size,
                        "ext": ext,
                        "is_binary": ext in _BINARY_EXTENSIONS,
                        "is_image": ext in {
                            ".png", ".jpg", ".jpeg", ".gif", ".webp", ".svg", ".bmp",
                        },
                        "is_video": ext in {".mp4", ".webm", ".ogg", ".mov", ".avi"},
                        "is_audio": ext in {".mp3", ".wav", ".flac", ".aac", ".m4a", ".opus"},
                        "is_spreadsheet": ext in {".xlsx", ".xls", ".xlsm", ".ods"},
                        "is_document": ext in {".docx", ".doc"},
                    })

            return {
                "name": root.name if rel_prefix == "" else path.name,
                "path": rel_prefix or ".",
                "type": "directory",
                "children": children,
            }

        tree = _build(root)
        tree["available"] = True
        return tree

    # ------------------------------------------------------------------
    # Search
    # ------------------------------------------------------------------

    def search(self, source: str, query: str) -> list[dict[str, Any]]:
        """Search files by name or path."""
        q = query.lower().strip()
        if not q or len(q) < 2:
            return []

        try:
            _obj, root, _ = self._resolve_source(source)
        except FileManagerError:
            return []

        results: list[dict[str, Any]] = []
        for f in sorted(root.rglob("*")):
            if not f.is_file() or f.name.startswith("."):
                continue
            rel = str(f.relative_to(root))
            if q in f.name.lower() or q in rel.lower():
                ext = f.suffix.lower()
                results.append({
                    "name": f.name,
                    "path": rel,
                    "size_bytes": f.stat().st_size,
                    "ext": ext,
                    "is_binary": ext in _BINARY_EXTENSIONS,
                })
        return results

    # ------------------------------------------------------------------
    # Statistics
    # ------------------------------------------------------------------

    def get_statistics(self) -> dict[str, Any]:
        """Return combined statistics for both sources."""
        stats: dict[str, Any] = {}
        for src_name in ("research", "docs"):
            try:
                _obj, root, _ = self._resolve_source(src_name)
                total = 0
                total_size = 0
                for f in root.rglob("*"):
                    if f.is_file() and not f.name.startswith("."):
                        total += 1
                        total_size += f.stat().st_size
                stats[src_name] = {
                    "available": True,
                    "total_files": total,
                    "total_size_bytes": total_size,
                }
            except FileManagerError:
                stats[src_name] = {
                    "available": False,
                    "total_files": 0,
                    "total_size_bytes": 0,
                }
        return {"sources": stats}

    # ------------------------------------------------------------------
    # Content retrieval
    # ------------------------------------------------------------------

    def get_content(
        self, source: str, file_path: str
    ) -> tuple[bytes | str, str | None, bool]:
        """Get file content.

        Returns:
            Tuple of (content, mime_type_or_None, is_binary).
            For text files: (str, None, False).
            For binary files: (bytes, mime_type, True).
        """
        if not file_path or ".." in file_path:
            raise PathTraversalError("Path traversal is not allowed.")

        try:
            _obj, root, _ = self._resolve_source(source)
        except FileManagerError:
            raise SourceNotAvailableError(f"Source '{source}' is not available.")

        candidate = (root / file_path).resolve()
        try:
            candidate.relative_to(root)
        except ValueError:
            raise PathTraversalError("Path traversal detected.")

        if not candidate.is_file():
            raise FileNotFoundError(f"File not found: {file_path}")

        ext = candidate.suffix.lower()

        if ext in _BINARY_EXTENSIONS:
            data = candidate.read_bytes()
            return data, _mime_type(ext), True

        content = candidate.read_text(encoding="utf-8", errors="replace")
        return content, None, False

    def _resolve_file(self, source: str, file_path: str) -> Path:
        """Resolve and validate a file path under a source root."""
        if not file_path or ".." in file_path:
            raise PathTraversalError("Path traversal is not allowed.")

        try:
            _obj, root, _ = self._resolve_source(source)
        except FileManagerError:
            raise SourceNotAvailableError(f"Source '{source}' is not available.")

        candidate = (root / file_path).resolve()
        try:
            candidate.relative_to(root)
        except ValueError:
            raise PathTraversalError("Path traversal detected.")
        return candidate

    def save_content(
        self, source: str, file_path: str, content: str, backup: bool = True
    ) -> dict[str, Any]:
        """Save text content to a file under the configured source root.

        Args:
            source: 'research' or 'docs'.
            file_path: Relative path inside the source root.
            content: New file content as UTF-8 text.
            backup: If True, the existing file is renamed to `<name>.bak`
                before writing.

        Returns:
            Dict with success status and path information.
        """
        candidate = self._resolve_file(source, file_path)

        # Only allow saving to existing files for safety. New files can be
        # created explicitly later if needed.
        if not candidate.is_file():
            raise FileNotFoundError(f"File not found: {file_path}")

        ext = candidate.suffix.lower()
        if ext in _BINARY_EXTENSIONS:
            raise FileManagerError("Binary files cannot be edited as text.")

        if backup and candidate.exists():
            backup_path = candidate.with_suffix(candidate.suffix + ".bak")
            # Remove stale backup to avoid collision
            if backup_path.exists():
                backup_path.unlink()
            candidate.rename(backup_path)

        candidate.write_text(content, encoding="utf-8")
        return {
            "success": True,
            "path": file_path,
            "size_bytes": candidate.stat().st_size,
            "backup_created": backup,
        }

    def _meta_path(self, source: str, file_path: str) -> Path:
        """Return the sidecar metadata path for a file."""
        candidate = self._resolve_file(source, file_path)
        return candidate.parent / (candidate.name + ".meta.yaml")

    def get_meta(self, source: str, file_path: str) -> dict[str, Any]:
        """Load sidecar metadata for a file.

        Args:
            source: 'research' or 'docs'.
            file_path: Relative path inside the source root.

        Returns:
            Dict with meta content and existence status.
        """
        candidate = self._resolve_file(source, file_path)
        if not candidate.is_file():
            raise FileNotFoundError(f"File not found: {file_path}")

        meta_path = self._meta_path(source, file_path)
        content = meta_path.read_text(encoding="utf-8") if meta_path.is_file() else ""
        return {
            "path": file_path,
            "meta_path": str(meta_path.relative_to(self._root(source))),
            "exists": meta_path.is_file(),
            "content": content,
        }

    def save_meta(
        self, source: str, file_path: str, content: str, backup: bool = True
    ) -> dict[str, Any]:
        """Save sidecar metadata for a file.

        Args:
            source: 'research' or 'docs'.
            file_path: Relative path inside the source root.
            content: YAML metadata text.
            backup: If True, existing meta file is backed up.

        Returns:
            Dict with success status and path information.
        """
        candidate = self._resolve_file(source, file_path)
        if not candidate.is_file():
            raise FileNotFoundError(f"File not found: {file_path}")

        meta_path = self._meta_path(source, file_path)
        if backup and meta_path.exists():
            backup_path = meta_path.with_suffix(".meta.yaml.bak")
            if backup_path.exists():
                backup_path.unlink()
            meta_path.rename(backup_path)

        meta_path.write_text(content, encoding="utf-8")
        return {
            "success": True,
            "path": file_path,
            "meta_path": str(meta_path.relative_to(self._root(source))),
            "size_bytes": meta_path.stat().st_size,
            "backup_created": backup,
        }

    def analyze_content(self, source: str, file_path: str) -> dict[str, Any]:
        """Return a lightweight local analysis of a text file.

        Args:
            source: 'research' or 'docs'.
            file_path: Relative path inside the source root.

        Returns:
            Dict with statistics, language guess, readability, keywords,
            sentiment and a generated summary.
        """
        candidate = self._resolve_file(source, file_path)
        if not candidate.is_file():
            raise FileNotFoundError(f"File not found: {file_path}")

        ext = candidate.suffix.lower()
        if ext in _BINARY_EXTENSIONS:
            raise FileManagerError("Binary files cannot be analyzed as text.")

        content = candidate.read_text(encoding="utf-8")
        if not content.strip():
            return {
                "path": file_path,
                "language": None,
                "stats": {"chars": 0, "words": 0, "lines": 0, "sentences": 0},
                "readability": None,
                "keywords": [],
                "sentiment": {"score": 0, "label": "neutral"},
                "summary": "",
            }

        import re

        # Basic tokenization
        text_lower = content.lower()
        words = re.findall(r"\b[a-zA-ZäöüßÄÖÜ]{3,}\b", text_lower)
        sentences = re.split(r"(?<=[.!?])\s+", content.strip())
        sentences = [s for s in sentences if s.strip()]

        # Language guess
        de_markers = sum(1 for w in ["der", "die", "das", "und", "ist", "von", "mit", "für", "den", "dem"] if w in words)
        en_markers = sum(1 for w in ["the", "and", "is", "of", "to", "a", "in", "for", "that", "with"] if w in words)
        language = "de" if de_markers > en_markers else "en"

        # Stopwords
        stopwords = {
            "en": {"the", "and", "is", "of", "to", "a", "in", "for", "that", "with", "as", "on", "by", "this", "from", "it", "be", "are", "was", "were", "been", "have", "has", "had", "do", "does", "did", "will", "would", "could", "should", "may", "might", "can", "shall"},
            "de": {"der", "die", "das", "und", "ist", "von", "mit", "für", "den", "dem", "ein", "eine", "einer", "eines", "einem", "einen", "zu", "auf", "an", "in", "bei", "nach", "aus", "durch", "wie", "so", "wenn", "dann", "als", "auch", "noch", "nur", "oder", "aber", "sondern", "weil", "dass"},
        }[language]

        # Keywords by frequency
        filtered = [w for w in words if w not in stopwords]
        counts: dict[str, int] = {}
        for w in filtered:
            counts[w] = counts.get(w, 0) + 1
        keywords = sorted(counts.items(), key=lambda x: x[1], reverse=True)[:10]

        # Readability (simplified)
        word_count = len(words)
        sentence_count = max(len(sentences), 1)
        avg_sentence_length = word_count / sentence_count
        if language == "en":
            # Flesch Reading Ease approximation
            syllables = sum(max(1, len(re.findall(r"[aeiouy]+", w))) for w in words)
            asl = avg_sentence_length
            asw = syllables / max(word_count, 1)
            flesch = 206.835 - (1.015 * asl) - (84.6 * asw)
            readability = {"score": round(flesch, 2), "label": "flesch_reading_ease"}
        else:
            readability = {"score": round(avg_sentence_length, 2), "label": "avg_words_per_sentence"}

        # Sentiment
        positive = {"good", "great", "excellent", "positive", "success", "improve", "benefit", "best", "happy", "love", "nice", "easy", "strong", "gut", "grossartig", "ausgezeichnet", "positiv", "erfolg", "verbessern", "vorteil", "beste", "glücklich", "liebe", "schön", "einfach", "stark"}
        negative = {"bad", "error", "fail", "problem", "bug", "issue", "wrong", "poor", "negative", "difficult", "schlecht", "fehler", "scheitern", "problem", "bug", "problem", "falsch", "arm", "negativ", "schwierig"}
        pos_count = sum(1 for w in words if w in positive)
        neg_count = sum(1 for w in words if w in negative)
        sentiment_score = pos_count - neg_count
        if sentiment_score > 0:
            sentiment_label = "positive"
        elif sentiment_score < 0:
            sentiment_label = "negative"
        else:
            sentiment_label = "neutral"

        # Summary: first few sentences, capped
        summary_sentences = sentences[:3]
        summary = " ".join(summary_sentences)

        return {
            "path": file_path,
            "language": language,
            "stats": {
                "chars": len(content),
                "words": word_count,
                "lines": content.count("\n") + 1,
                "sentences": len(sentences),
            },
            "readability": readability,
            "keywords": [{"word": w, "count": c} for w, c in keywords],
            "sentiment": {"score": sentiment_score, "label": sentiment_label},
            "summary": summary,
        }

    def export_content(
        self, source: str, file_path: str, fmt: str
    ) -> tuple[str, bytes, str]:
        """Export a text file to HTML, DOCX or Markdown.

        Args:
            source: 'research' or 'docs'.
            file_path: Relative path inside the source root.
            fmt: 'html', 'docx' or 'md'.

        Returns:
            Tuple of (filename, bytes, mime_type).
        """
        candidate = self._resolve_file(source, file_path)
        if not candidate.is_file():
            raise FileNotFoundError(f"File not found: {file_path}")

        ext = candidate.suffix.lower()
        if ext in _BINARY_EXTENSIONS:
            raise FileManagerError("Binary files cannot be exported as documents.")

        content = candidate.read_text(encoding="utf-8")
        base_name = candidate.stem
        fmt = fmt.lower()

        if fmt == "md":
            return (
                f"{base_name}.md",
                content.encode("utf-8"),
                "text/markdown; charset=utf-8",
            )

        if fmt == "html":
            html = _markdown_to_html(content, title=base_name)
            return (
                f"{base_name}.html",
                html.encode("utf-8"),
                "text/html; charset=utf-8",
            )

        if fmt == "docx":
            try:
                from docx import Document
            except ImportError as exc:  # pragma: no cover - optional dependency
                raise FileManagerError(
                    "python-docx is not installed. Install it to enable DOCX export."
                ) from exc

            doc = Document()
            doc.add_heading(base_name, level=0)
            for block in _iter_markdown_blocks(content):
                if block["type"] == "heading":
                    doc.add_heading(block["text"], level=block["level"])
                elif block["type"] == "paragraph":
                    doc.add_paragraph(block["text"])
                elif block["type"] == "code":
                    doc.add_paragraph(block["text"], style="Quote")
                elif block["type"] == "list":
                    doc.add_paragraph(block["text"], style="List Bullet")
            out = BytesIO()
            doc.save(out)
            return (
                f"{base_name}.docx",
                out.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

        raise FileManagerError(f"Unsupported export format: {fmt}")

    def get_history(self, source: str, file_path: str, limit: int = 20) -> list[dict[str, Any]]:
        """Return git history for a file under the configured source root.

        Args:
            source: 'research' or 'docs'.
            file_path: Relative path inside the source root.
            limit: Maximum number of commits to return.

        Returns:
            List of commit dicts with hash, date, author, message.
        """
        candidate = self._resolve_file(source, file_path)
        if not candidate.is_file():
            raise FileNotFoundError(f"File not found: {file_path}")

        root = candidate.parent
        while not (root / ".git").is_dir() and root != root.parent:
            root = root.parent

        if not (root / ".git").is_dir():
            return []

        try:
            rel = candidate.relative_to(root)
        except ValueError:
            return []

        try:
            result = subprocess.run(
                [
                    "git", "log", f"-n{limit}", "--follow", "--oneline",
                    "--format=%H|%ad|%an|%s", "--date=iso-strict", "--", str(rel)
                ],
                cwd=str(root),
                capture_output=True,
                text=True,
                encoding="utf-8",
                errors="replace",
                check=False,
            )
        except FileNotFoundError:
            return []

        if result.returncode != 0:
            return []

        history: list[dict[str, str]] = []
        for line in result.stdout.strip().splitlines():
            parts = line.split("|", 3)
            if len(parts) < 4:
                continue
            history.append({
                "hash": parts[0],
                "date": parts[1],
                "author": parts[2],
                "message": parts[3],
            })
        return history


# ---------------------------------------------------------------------------
# HTTP handler integration
# ---------------------------------------------------------------------------

_DEFAULT_DOCS_ROOT = Path(__file__).resolve().parents[2] / "docs"


def register_file_manager_routes(
    handler: Any,
    path: str,
    query: dict[str, list[str]],
    research_source: ResearchSource | None,
    docs_source: DocumentationSource | None,
) -> bool:
    """Try to handle a file manager API route. Returns True if handled."""
    fm = FileManager(research_source, docs_source, _DEFAULT_DOCS_ROOT)

    if path == "/api/files/tree":
        source = query.get("source", ["research"])[0]
        handler._send_json(fm.get_tree(source))
        return True

    if path == "/api/files/search":
        source = query.get("source", ["research"])[0]
        q = query.get("q", [""])[0]
        handler._send_json({"results": fm.search(source, q)})
        return True

    if path == "/api/files/statistics":
        handler._send_json(fm.get_statistics())
        return True

    if path == "/api/files/open":
        source = query.get("source", ["research"])[0]
        try:
            _obj, root, _name = fm._resolve_source(source)  # pyright: ignore[reportPrivateUsage]
            import os
            import subprocess
            if os.name == "nt":
                subprocess.Popen(["explorer", str(root.resolve())])
            else:
                subprocess.Popen(["xdg-open", str(root.resolve())])
            handler._send_json({"ok": True, "path": str(root.resolve())})
        except Exception as exc:
            handler._send_json({"ok": False, "error": str(exc)})
        return True

    if path.startswith("/api/files/content/"):
        prefix = "/api/files/content/"
        file_path = unquote(path[len(prefix):])
        source = query.get("source", ["research"])[0]

        try:
            content, mime, is_binary = fm.get_content(source, file_path)
        except FileNotFoundError:
            handler._send_json(
                {"error": f"File not found: {file_path}"},
                HTTPStatus.NOT_FOUND,
            )
            return True
        except (SourceNotAvailableError, PathTraversalError) as exc:
            status = HTTPStatus.NOT_FOUND if isinstance(exc, SourceNotAvailableError) else HTTPStatus.FORBIDDEN
            handler._send_json({"error": str(exc)}, status)
            return True
        except OSError as exc:
            handler._send_json(
                {"error": str(exc)},
                HTTPStatus.INTERNAL_SERVER_ERROR,
            )
            return True

        if is_binary and mime:
            data = cast(bytes, content)
            handler.send_response(HTTPStatus.OK)
            handler.send_header("Content-Type", mime)
            handler.send_header("Content-Length", str(len(data)))
            handler.send_header("Cache-Control", "no-cache")
            handler.send_header("Access-Control-Allow-Origin", "*")
            handler.end_headers()
            handler.wfile.write(data)
        else:
            text_content = cast(str, content)
            handler._send_json({
                "path": file_path,
                "name": Path(file_path).name,
                "content": text_content,
                "size_bytes": len(text_content.encode("utf-8")),
                "ext": Path(file_path).suffix.lower(),
                "is_binary": False,
            })
        return True

    if path.startswith("/api/files/save/") and handler.command == "PUT":
        prefix = "/api/files/save/"
        file_path = unquote(path[len(prefix):])
        source = query.get("source", ["research"])[0]
        try:
            body = handler._read_json_body()
        except Exception:
            handler._send_json(
                {"error": "Invalid JSON body"},
                HTTPStatus.BAD_REQUEST,
            )
            return True

        content = body.get("content")
        if content is None or not isinstance(content, str):
            handler._send_json(
                {"error": "Missing or invalid 'content' field"},
                HTTPStatus.BAD_REQUEST,
            )
            return True

        try:
            result = fm.save_content(source, file_path, content, backup=body.get("backup", True))
            handler._send_json(result)
        except FileNotFoundError as exc:
            handler._send_json({"error": str(exc)}, HTTPStatus.NOT_FOUND)
        except PathTraversalError as exc:
            handler._send_json({"error": str(exc)}, HTTPStatus.FORBIDDEN)
        except FileManagerError as exc:
            handler._send_json({"error": str(exc)}, HTTPStatus.BAD_REQUEST)
        except OSError as exc:
            handler._send_json({"error": str(exc)}, HTTPStatus.INTERNAL_SERVER_ERROR)
        return True

    if path.startswith("/api/files/history/"):
        prefix = "/api/files/history/"
        file_path = unquote(path[len(prefix):])
        source = query.get("source", ["research"])[0]
        limit = int(query.get("limit", ["20"])[0])

        try:
            history = fm.get_history(source, file_path, limit=limit)
            handler._send_json({"path": file_path, "history": history})
        except FileNotFoundError as exc:
            handler._send_json({"error": str(exc)}, HTTPStatus.NOT_FOUND)
        except PathTraversalError as exc:
            handler._send_json({"error": str(exc)}, HTTPStatus.FORBIDDEN)
        return True

    if path.startswith("/api/files/meta/") and handler.command == "GET":
        prefix = "/api/files/meta/"
        file_path = unquote(path[len(prefix):])
        source = query.get("source", ["research"])[0]

        try:
            result = fm.get_meta(source, file_path)
            handler._send_json(result)
        except FileNotFoundError as exc:
            handler._send_json({"error": str(exc)}, HTTPStatus.NOT_FOUND)
        except PathTraversalError as exc:
            handler._send_json({"error": str(exc)}, HTTPStatus.FORBIDDEN)
        return True

    if path.startswith("/api/files/meta/") and handler.command == "PUT":
        prefix = "/api/files/meta/"
        file_path = unquote(path[len(prefix):])
        source = query.get("source", ["research"])[0]
        try:
            body = handler._read_json_body()
        except Exception:
            handler._send_json(
                {"error": "Invalid JSON body"},
                HTTPStatus.BAD_REQUEST,
            )
            return True

        content = body.get("content")
        if content is None or not isinstance(content, str):
            handler._send_json(
                {"error": "Missing or invalid 'content' field"},
                HTTPStatus.BAD_REQUEST,
            )
            return True

        try:
            result = fm.save_meta(source, file_path, content, backup=body.get("backup", True))
            handler._send_json(result)
        except FileNotFoundError as exc:
            handler._send_json({"error": str(exc)}, HTTPStatus.NOT_FOUND)
        except PathTraversalError as exc:
            handler._send_json({"error": str(exc)}, HTTPStatus.FORBIDDEN)
        except FileManagerError as exc:
            handler._send_json({"error": str(exc)}, HTTPStatus.BAD_REQUEST)
        except OSError as exc:
            handler._send_json({"error": str(exc)}, HTTPStatus.INTERNAL_SERVER_ERROR)
        return True

    if path.startswith("/api/files/analyze/"):
        prefix = "/api/files/analyze/"
        file_path = unquote(path[len(prefix):])
        source = query.get("source", ["research"])[0]

        try:
            result = fm.analyze_content(source, file_path)
            handler._send_json(result)
        except FileNotFoundError as exc:
            handler._send_json({"error": str(exc)}, HTTPStatus.NOT_FOUND)
        except PathTraversalError as exc:
            handler._send_json({"error": str(exc)}, HTTPStatus.FORBIDDEN)
        except FileManagerError as exc:
            handler._send_json({"error": str(exc)}, HTTPStatus.BAD_REQUEST)
        return True

    if path.startswith("/api/files/export/"):
        prefix = "/api/files/export/"
        file_path = unquote(path[len(prefix):])
        source = query.get("source", ["research"])[0]
        fmt = query.get("format", ["html"])[0]

        try:
            filename, data, mime = fm.export_content(source, file_path, fmt)
            handler.send_response(HTTPStatus.OK)
            handler.send_header("Content-Type", mime)
            handler.send_header("Content-Disposition", f'attachment; filename="{filename}"')
            handler.send_header("Content-Length", str(len(data)))
            handler.send_header("Cache-Control", "no-cache")
            handler.send_header("Access-Control-Allow-Origin", "*")
            handler.end_headers()
            handler.wfile.write(data)
        except FileNotFoundError as exc:
            handler._send_json({"error": str(exc)}, HTTPStatus.NOT_FOUND)
        except PathTraversalError as exc:
            handler._send_json({"error": str(exc)}, HTTPStatus.FORBIDDEN)
        except FileManagerError as exc:
            handler._send_json({"error": str(exc)}, HTTPStatus.BAD_REQUEST)
        return True

    return False
